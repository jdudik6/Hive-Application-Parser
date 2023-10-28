import csv
import docx
import openpyxl

# Resources Utilizied:
# https://docs.python.org/3/library/csv.html
# https://www.geeksforgeeks.org/python-working-with-docx-module/
# https://www.geeksforgeeks.org/python-reading-excel-file-using-openpyxl-module/


doc = docx.Document()
title = doc.add_heading('Spring 2024 Applications', 0)
title.alignment = 1
doc.add_paragraph('Use CTRL + F to search for the heading containing your applicant''s name.\nUse the forms below to submit your scores and comments.')
doc.add_paragraph('Scores: https://airtable.com/app3NdNkOfXRNjmSF/shrH56epcGGHBrKjR')
doc.add_paragraph('Comments: https://airtable.com/app3NdNkOfXRNjmSF/shrvAw2YR3q5j2mAE')
doc.add_page_break()

wb_xlsx = openpyxl.load_workbook('Applications.xlsx')
app_xlsx = wb_xlsx.active

num_applicants = app_xlsx.max_row

print('Applicants: ' + str(num_applicants - 1))


for applicant_num in range(1, num_applicants):
    for field in range(7, 17):
        field_val = str(app_xlsx.cell(row = applicant_num + 1, column = field).value)

        match field:
            case 7:
                full_name = field_val + ' ' + str(app_xlsx.cell(row = applicant_num + 1, column = field + 1).value)
                doc.add_heading(full_name, 1)

            case 9:
                if (field_val != 'None'):
                    doc.add_heading('Preferred Name: ' + field_val, 2)
            
            case 10:
                doc.add_heading('Current Year: ' + field_val, 3)
                
            case 11:
                doc.add_heading('Major: ' + field_val, 3)
                
            case 12:
                doc.add_heading('Why do you want to be a Peer Instructor at the Hive?', 3)
                doc.add_paragraph(field_val)
                
            case 14:
                doc.add_heading('What can you bring to the Hive?', 3)
                doc.add_paragraph(field_val)
                
            case 15:
                doc.add_heading('"Describe any current and previous extracurricular activities you are involved in, including both on-campus and off-campus. If applicable, include any teaching and volunteer experience.', 3)
                doc.add_paragraph(field_val)
                
            case 16:
                doc.add_heading('Keep in mind that as an active Peer Instructor, you are required to be on shift for a total of 3 hours per week. List your current and projected commitments and/or workload, along with the hours you spend weekly on each.', 3)
                doc.add_paragraph(field_val)

            case 17:
                doc.add_heading('Do you understand the responsibilities of being a Peer Instructor at the Hive, including, but not limited to: completion of trainings within the training period, upkeep of the makerspace, and the equired 3 hours per week on-shift?', 3)
                doc.add_paragraph(field_val)

            case 18:
                if (field_val != 'None'):
                    doc.add_heading('Anything you would like us to consider?', 3)
                    doc.add_paragraph(field_val)

            case _:
                continue

    doc.add_page_break()
            
doc.save('Applications.docx')
print('Applications processed successfully.')
input("Press enter to exit:")
    

